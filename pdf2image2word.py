'''
pip3 install pdf2image
pip3 install docx
pip3 install python-docx 
'''

from pdf2image import convert_from_path

from docx import Document
from docx.shared import Inches

# Open new word document
document = Document()

# Set page to A4 size and change the page margins (use inches).
# Set margin to 0
sections = document.sections
for section in sections:
	section.top_margin 		= Inches(0)
	section.bottom_margin 	= Inches(0)
	section.left_margin 	= Inches(0)
	section.right_margin 	= Inches(0)
	
	section.page_height 	= Inches(11.69)
	section.page_width 		= Inches(8.27)

p = document.add_paragraph()
r = p.add_run()

# Path to pdf, set the image quality to 500 
pages = convert_from_path('/Users/lmkhuyen/Desktop/file.pdf', 500)

# Export each pdf page to image, then insert it to word file
i = 1
for page in pages:
	image_name = 'out' + str(i) + '.jpg'
	page.save(image_name, 'JPEG')
	r.add_picture(image_name, width=Inches(8.27))  # Re-size image width to 8.27 inches, so it will fit the image 
	i = i + 1

document.save('lmkhuyen_word.docx')
